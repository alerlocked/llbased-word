"""
文档导出API
支持导出Word、PDF等格式
"""
import io
import os
import tempfile
from pathlib import Path

from fastapi import APIRouter, HTTPException, Depends
from fastapi.responses import FileResponse, StreamingResponse
from sqlalchemy.orm import Session
from pydantic import BaseModel
from typing import List, Optional

from app.database import get_db
from app.models.database import CreationProject, Annotation, Citation, ArticleFigure
from app.services.word_export_service import get_word_export_service
from app.shared.logging import get_logger
logger = get_logger(__name__)
from app.utils.path_utils import build_static_url

router = APIRouter()


class ExportWordRequest(BaseModel):
    """Word导出请求"""
    project_id: int
    include_annotations: bool = True
    include_citations: bool = True
    include_figures: bool = True


@router.post("/word")
async def export_word(
    request: ExportWordRequest,
    db: Session = Depends(get_db)
):
    """
    导出项目为Word文档
    
    Args:
        project_id: 项目ID
        include_annotations: 是否包含注释
        include_citations: 是否包含引用
        include_figures: 是否包含图片
    """
    logger.info(f"📤 API: 导出Word文档 - 项目ID:{request.project_id}")
    
    try:
        # 查找项目
        project = db.query(CreationProject).filter(
            CreationProject.id == request.project_id
        ).first()
        
        if not project:
            raise HTTPException(status_code=404, detail="项目不存在")
        
        if not project.content:
            raise HTTPException(status_code=400, detail="项目内容为空")
        
        # 获取注释
        annotations = []
        if request.include_annotations:
            ann_records = db.query(Annotation).filter(
                Annotation.project_id == request.project_id
            ).all()
            
            annotations = [
                {
                    "content": ann.content,
                    "annotation_type": ann.annotation_type,
                    "position": ann.position
                }
                for ann in ann_records
            ]
        
        # 获取引用
        citations = []
        if request.include_citations:
            cit_records = db.query(Citation).filter(
                Citation.project_id == request.project_id
            ).order_by(Citation.citation_number).all()
            
            citations = [
                {
                    "source_type": cit.source_type,
                    "author": cit.author,
                    "title": cit.title,
                    "date": cit.date,
                    "url": cit.url,
                    "content": cit.content
                }
                for cit in cit_records
            ]
        
        # 获取图片
        figures = []
        if request.include_figures:
            fig_records = db.query(ArticleFigure).filter(
                ArticleFigure.project_id == request.project_id
            ).order_by(ArticleFigure.figure_number).all()
            
            figures = [
                {
                    "figure_number": fig.figure_number,
                    "caption": fig.caption,
                    "position": fig.position
                }
                for fig in fig_records
            ]
        
        # 预处理内容：查找并替换网络图片URL为本地路径
        # 从内容中提取Markdown图片语法，查找已下载的本地路径
        processed_content = project.content
        if processed_content:
            import re
            from app.models.database import WebImage
            
            # 查找所有Markdown图片语法
            markdown_images = re.finditer(r'!\[(.*?)\]\((.*?)\)', processed_content)
            replacements = []
            
            for match in markdown_images:
                image_url = match.group(2).strip()
                # 如果是网络URL，查找本地路径
                if image_url.startswith('http://') or image_url.startswith('https://'):
                    web_image = db.query(WebImage).filter(
                        WebImage.original_url == image_url
                    ).first()
                    
                    if web_image and web_image.local_path:
                        # 替换为本地路径（相对路径格式）
                        local_path = build_static_url(web_image.local_path)
                        replacements.append((match.group(0), f"![{match.group(1)}]({local_path})"))
            
            # 执行替换
            for old_text, new_text in replacements:
                processed_content = processed_content.replace(old_text, new_text)
        
        # 导出Word
        word_service = get_word_export_service()
        file_path = word_service.export_article(
            title=project.name,
            content=processed_content,
            annotations=annotations if annotations else None,
            citations=citations if citations else None,
            figures=figures if figures else None,
            filename=f"{project.name}_{project.id}"
        )
        
        logger.info(f"✅ Word导出成功: {file_path}")
        
        # 返回文件
        return FileResponse(
            path=str(file_path),
            filename=f"{project.name}.docx",
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"❌ Word导出失败: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))


@router.get("/word/preview/{project_id}")
async def preview_word_content(
    project_id: int,
    db: Session = Depends(get_db)
):
    """
    预览Word导出内容（不实际导出）
    
    Args:
        project_id: 项目ID
    """
    logger.info(f"👁️ API: 预览Word内容 - 项目ID:{project_id}")
    
    try:
        # 查找项目
        project = db.query(CreationProject).filter(
            CreationProject.id == project_id
        ).first()
        
        if not project:
            raise HTTPException(status_code=404, detail="项目不存在")
        
        # 获取各种内容的数量
        annotation_count = db.query(Annotation).filter(
            Annotation.project_id == project_id
        ).count()
        
        citation_count = db.query(Citation).filter(
            Citation.project_id == project_id
        ).count()
        
        figure_count = db.query(ArticleFigure).filter(
            ArticleFigure.project_id == project_id
        ).count()
        
        return {
            "project_id": project.id,
            "title": project.name,
            "content_length": len(project.content) if project.content else 0,
            "annotation_count": annotation_count,
            "citation_count": citation_count,
            "figure_count": figure_count,
            "can_export": project.content is not None and len(project.content) > 0
        }

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"❌ 预览失败: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))


# ── Content-based export (no project_id required) ──────────────

class ExportContentRequest(BaseModel):
    """Content-based export request — accepts raw content"""
    title: str = "未命名文档"
    content: str


@router.post("/content-pdf")
async def export_content_pdf(request: ExportContentRequest):
    """Export raw markdown/text content as PDF (no project required)."""
    logger.info(f"📤 API: 内容PDF导出 - title={request.title}, chars={len(request.content)}")
    if not request.content.strip():
        raise HTTPException(status_code=400, detail="内容为空")
    try:
        pdf_bytes = _generate_pdf_from_content(request.title, request.content)
        safe_title = request.title.replace("/", "_").replace("\\", "_")
        return StreamingResponse(
            io.BytesIO(pdf_bytes),
            media_type="application/pdf",
            headers={"Content-Disposition": f'attachment; filename="{safe_title}.pdf"'},
        )
    except Exception as e:
        logger.error(f"❌ 内容PDF导出失败: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/content-word")
async def export_content_word(request: ExportContentRequest):
    """Export raw markdown/text content as Word (no project required)."""
    logger.info(f"📤 API: 内容Word导出 - title={request.title}, chars={len(request.content)}")
    if not request.content.strip():
        raise HTTPException(status_code=400, detail="内容为空")
    try:
        word_path = _generate_word_from_content(request.title, request.content)
        safe_title = request.title.replace("/", "_").replace("\\", "_")
        return FileResponse(
            path=str(word_path),
            filename=f"{safe_title}.docx",
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
    except Exception as e:
        logger.error(f"❌ 内容Word导出失败: {e}")
        raise HTTPException(status_code=500, detail=str(e))


def _generate_pdf_from_content(title: str, content: str) -> bytes:
    """Generate PDF from text content using reportlab."""
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import mm
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    import re

    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf, pagesize=A4,
        leftMargin=25 * mm, rightMargin=25 * mm,
        topMargin=20 * mm, bottomMargin=20 * mm,
    )

    styles = getSampleStyleSheet()

    # Register CJK font
    font_name = "Helvetica"
    for font_path in [
        "C:/Windows/Fonts/msyh.ttc",
        "C:/Windows/Fonts/simsun.ttc",
        "C:/Windows/Fonts/simhei.ttf",
    ]:
        if os.path.exists(font_path):
            try:
                pdfmetrics.registerFont(TTFont("CJKFont", font_path))
                font_name = "CJKFont"
                break
            except Exception:
                continue

    title_style = ParagraphStyle("ExportTitle", parent=styles["Title"], fontName=font_name, fontSize=18, spaceAfter=12)
    h2_style = ParagraphStyle("ExportH2", parent=styles["Heading2"], fontName=font_name, fontSize=14, spaceBefore=12, spaceAfter=6)
    body_style = ParagraphStyle("ExportBody", parent=styles["Normal"], fontName=font_name, fontSize=11, leading=16, spaceBefore=2)

    def _escape(text: str) -> str:
        return text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")

    elements = []
    elements.append(Paragraph(_escape(title), title_style))
    elements.append(Spacer(1, 12))

    for line in content.splitlines():
        stripped = line.strip()
        if not stripped:
            elements.append(Spacer(1, 4))
            continue
        # Detect markdown headings
        if stripped.startswith("## "):
            elements.append(Paragraph(_escape(stripped[3:]), h2_style))
        elif stripped.startswith("# "):
            elements.append(Paragraph(_escape(stripped[2:]), title_style))
        else:
            elements.append(Paragraph(_escape(stripped), body_style))

    doc.build(elements)
    return buf.getvalue()


def _generate_word_from_content(title: str, content: str) -> Path:
    """Generate Word from text content using python-docx."""
    from docx import Document
    from docx.shared import Pt

    doc = Document()
    doc.add_heading(title, level=0)

    for line in content.splitlines():
        stripped = line.strip()
        if not stripped:
            doc.add_paragraph("")
            continue
        if stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith("# "):
            doc.add_heading(stripped[2:], level=1)
        elif stripped.startswith("| ") and "|" in stripped[2:]:
            # Table row — collect consecutive table lines
            para = doc.add_paragraph(stripped)
            para.style.font.size = Pt(10)
        else:
            para = doc.add_paragraph(stripped)
            para.style.font.size = Pt(11)

    tmp_dir = Path(tempfile.gettempdir()) / "content_exports"
    tmp_dir.mkdir(parents=True, exist_ok=True)
    safe_title = title.replace("/", "_").replace("\\", "_")
    file_path = tmp_dir / f"{safe_title}.docx"
    doc.save(str(file_path))
    return file_path

