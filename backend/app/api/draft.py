# -*- coding: utf-8 -*-
"""
Draft API - 工艺文件初稿管理端点
Phase 4 - PIV: piv_20260411_draft_api_phase4
"""
import io
import os
import shutil
import tempfile
from datetime import datetime
from pathlib import Path
from typing import Optional

from fastapi import APIRouter, Depends, HTTPException, UploadFile, File, Query
from fastapi.responses import FileResponse, StreamingResponse
from pydantic import BaseModel
from sqlalchemy.orm import Session

from app.database import get_db
from app.models.database import DraftDocument, DraftVersion
from app.services.draft_service import DraftService
from app.shared.logging import get_logger

logger = get_logger(__name__)

router = APIRouter()


# ── Response models ──────────────────────────────────────────────

class DraftResponse(BaseModel):
    id: int
    title: str
    file_type: Optional[str] = None
    status: Optional[str] = None
    project_id: Optional[int] = None
    content: Optional[str] = None
    created_at: Optional[datetime] = None
    updated_at: Optional[datetime] = None

    class Config:
        from_attributes = True


class VersionResponse(BaseModel):
    id: int
    draft_id: int
    snapshot_source: Optional[str] = None
    snapshot_content: Optional[str] = None
    created_at: Optional[datetime] = None

    class Config:
        from_attributes = True


# ── 1. Upload ────────────────────────────────────────────────────

@router.post("/upload", response_model=DraftResponse)
async def upload_draft(
    file: UploadFile = File(...),
    project_id: Optional[int] = Query(None),
    db: Session = Depends(get_db),
):
    """上传初稿文件（PDF/Word）"""
    logger.info(f"📤 Draft API: upload_draft - filename={file.filename}")
    try:
        svc = DraftService(db)
        draft = await svc.upload_draft(file, project_id=project_id)
        return draft
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))
    except Exception as e:
        logger.error(f"❌ 上传初稿失败: {e}")
        raise HTTPException(status_code=500, detail=str(e))


# ── 2. List drafts ──────────────────────────────────────────────

@router.get("", response_model=list[DraftResponse])
async def list_drafts(
    project_id: Optional[int] = Query(None),
    db: Session = Depends(get_db),
):
    """列出所有初稿，可选按 project_id 过滤"""
    logger.info(f"📋 Draft API: list_drafts - project_id={project_id}")
    svc = DraftService(db)
    return svc.list_drafts(project_id=project_id)


# ── 3. Get current content ──────────────────────────────────────

@router.get("/{draft_id}", response_model=DraftResponse)
async def get_draft(
    draft_id: int,
    db: Session = Depends(get_db),
):
    """获取初稿当前内容"""
    logger.info(f"📄 Draft API: get_draft - id={draft_id}")
    svc = DraftService(db)
    draft = svc.get_draft(draft_id)
    if draft is None:
        raise HTTPException(status_code=404, detail="初稿不存在")
    return draft


# ── 4. Version history ──────────────────────────────────────────

@router.get("/{draft_id}/versions", response_model=list[VersionResponse])
async def list_versions(
    draft_id: int,
    limit: int = Query(50, ge=1, le=200),
    db: Session = Depends(get_db),
):
    """版本历史列表（按时间倒序）"""
    logger.info(f"📚 Draft API: list_versions - draft_id={draft_id}, limit={limit}")
    svc = DraftService(db)
    draft = svc.get_draft(draft_id)
    if draft is None:
        raise HTTPException(status_code=404, detail="初稿不存在")
    return svc.list_versions(draft_id, limit=limit)


# ── 5. Get specific version ─────────────────────────────────────

@router.get("/{draft_id}/versions/{version_id}", response_model=VersionResponse)
async def get_version(
    draft_id: int,
    version_id: int,
    db: Session = Depends(get_db),
):
    """获取某版本内容"""
    logger.info(f"📄 Draft API: get_version - draft_id={draft_id}, version_id={version_id}")
    svc = DraftService(db)
    version = svc.get_version(version_id)
    if version is None or version.draft_id != draft_id:
        raise HTTPException(status_code=404, detail="版本不存在")
    return version


# ── 6. Rollback ─────────────────────────────────────────────────

@router.post("/{draft_id}/rollback/{version_id}", response_model=DraftResponse)
async def rollback_draft(
    draft_id: int,
    version_id: int,
    db: Session = Depends(get_db),
):
    """回滚到指定版本"""
    logger.info(f"⏪ Draft API: rollback - draft_id={draft_id}, version_id={version_id}")
    svc = DraftService(db)
    draft = svc.rollback(draft_id, version_id)
    if draft is None:
        raise HTTPException(status_code=404, detail="初稿或版本不存在")
    return draft


# ── 7. Diff ─────────────────────────────────────────────────────

@router.get("/{draft_id}/diff")
async def diff_versions(
    draft_id: int,
    v1: int = Query(..., description="版本1 ID"),
    v2: int = Query(..., description="版本2 ID"),
    db: Session = Depends(get_db),
):
    """两个版本的差异对比"""
    logger.info(f"🔍 Draft API: diff - draft_id={draft_id}, v1={v1}, v2={v2}")
    svc = DraftService(db)
    result = svc.get_diff(draft_id, v1, v2)
    if "error" in result:
        raise HTTPException(status_code=404, detail=result["error"])
    return result


# ── 8. Export PDF ───────────────────────────────────────────────

@router.post("/{draft_id}/export/pdf")
async def export_pdf(
    draft_id: int,
    db: Session = Depends(get_db),
):
    """导出初稿为 PDF 文件"""
    logger.info(f"📤 Draft API: export_pdf - draft_id={draft_id}")
    svc = DraftService(db)
    draft = svc.get_draft(draft_id)
    if draft is None:
        raise HTTPException(status_code=404, detail="初稿不存在")
    if not draft.content:
        raise HTTPException(status_code=400, detail="初稿内容为空，无法导出")

    try:
        pdf_bytes = _generate_pdf(draft.title, draft.content)
        filename = f"{draft.title or 'draft'}_{draft_id}.pdf"

        return StreamingResponse(
            io.BytesIO(pdf_bytes),
            media_type="application/pdf",
            headers={"Content-Disposition": f'attachment; filename="{filename}"'},
        )
    except Exception as e:
        logger.error(f"❌ PDF 导出失败: {e}")
        raise HTTPException(status_code=500, detail=f"PDF 导出失败: {e}")


# ── 9. Export Word ──────────────────────────────────────────────

@router.post("/{draft_id}/export/word")
async def export_word(
    draft_id: int,
    db: Session = Depends(get_db),
):
    """导出初稿为 Word 文件"""
    logger.info(f"📤 Draft API: export_word - draft_id={draft_id}")
    svc = DraftService(db)
    draft = svc.get_draft(draft_id)
    if draft is None:
        raise HTTPException(status_code=404, detail="初稿不存在")
    if not draft.content:
        raise HTTPException(status_code=400, detail="初稿内容为空，无法导出")

    try:
        word_path = _generate_word(draft.title, draft.content, draft_id)
        filename = f"{draft.title or 'draft'}_{draft_id}.docx"

        return FileResponse(
            path=str(word_path),
            filename=filename,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
    except Exception as e:
        logger.error(f"❌ Word 导出失败: {e}")
        raise HTTPException(status_code=500, detail=f"Word 导出失败: {e}")


# ── PDF generation helper ────────────────────────────────────────

def _generate_pdf(title: str, content: str) -> bytes:
    """使用 reportlab 将文本内容生成 PDF"""
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import mm
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont

    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf,
        pagesize=A4,
        leftMargin=25 * mm,
        rightMargin=25 * mm,
        topMargin=20 * mm,
        bottomMargin=20 * mm,
    )

    styles = getSampleStyleSheet()

    # Try to register a CJK font for Chinese content
    font_name = "Helvetica"
    try:
        # Try Windows system fonts
        for font_path in [
            "C:/Windows/Fonts/msyh.ttc",   # 微软雅黑
            "C:/Windows/Fonts/simsun.ttc",  # 宋体
            "C:/Windows/Fonts/simhei.ttf",  # 黑体
        ]:
            if os.path.exists(font_path):
                try:
                    pdfmetrics.registerFont(TTFont("CJKFont", font_path))
                    font_name = "CJKFont"
                    break
                except Exception:
                    continue
    except Exception:
        pass

    # Create styles with the available font
    title_style = ParagraphStyle(
        "DraftTitle",
        parent=styles["Title"],
        fontName=font_name,
        fontSize=18,
        spaceAfter=12,
    )
    body_style = ParagraphStyle(
        "DraftBody",
        parent=styles["Normal"],
        fontName=font_name,
        fontSize=11,
        leading=16,
        spaceBefore=4,
    )

    elements = []

    # Title
    safe_title = title or "未命名初稿"
    elements.append(Paragraph(safe_title, title_style))
    elements.append(Spacer(1, 12))

    # Body content - split by lines and create paragraphs
    for line in (content or "").splitlines():
        safe_line = _escape_xml(line) if line.strip() else "&nbsp;"
        elements.append(Paragraph(safe_line, body_style))

    doc.build(elements)
    return buf.getvalue()


def _escape_xml(text: str) -> str:
    """Escape XML special characters for reportlab Paragraph"""
    return (
        text.replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
    )


# ── Word generation helper ───────────────────────────────────────

def _generate_word(title: str, content: str, draft_id: int) -> Path:
    """使用 python-docx 生成 Word 文档"""
    from docx import Document
    from docx.shared import Pt, Inches

    doc = Document()

    # Title
    heading = doc.add_heading(title or "未命名初稿", level=0)

    # Content paragraphs
    for line in (content or "").splitlines():
        if line.strip():
            para = doc.add_paragraph(line)
            para.style.font.size = Pt(11)
        else:
            doc.add_paragraph("")

    # Save to temp file
    tmp_dir = Path(tempfile.gettempdir()) / "draft_exports"
    tmp_dir.mkdir(parents=True, exist_ok=True)
    file_path = tmp_dir / f"draft_{draft_id}_{title or 'export'}.docx"
    doc.save(str(file_path))

    return file_path


# ── 10. Temp upload (no DB, parse-only) ──────────────────────────

class TempUploadResponse(BaseModel):
    filename: str
    file_type: str
    content: str
    content_format: str = "html"   # "html" or "text"
    char_count: int
    page_count: Optional[int] = None


@router.post("/upload-temp", response_model=TempUploadResponse)
async def upload_temp(file: UploadFile = File(...)):
    """Upload a file for temporary AI context — no DB storage.

    PDF: VL Service (MinerU/Qwen) pipeline → HTML per page
    Word: python-docx extraction → HTML
    Results are not persisted to DB.
    """
    filename = file.filename or "unknown"
    ext = Path(filename).suffix.lower().lstrip(".")
    if ext not in ("pdf", "docx", "doc"):
        raise HTTPException(status_code=400, detail=f"Unsupported file type: {ext}")

    content_bytes = await file.read()
    file_type = "pdf" if ext == "pdf" else "docx"

    if file_type == "pdf":
        content, page_count = await _parse_pdf_with_vl(content_bytes)
    else:
        content, page_count = _parse_docx_to_html(content_bytes)

    logger.info(f"📤 Draft API: upload_temp - {filename}, {len(content)} chars, {page_count} pages")

    return TempUploadResponse(
        filename=filename,
        file_type=file_type,
        content=content,
        content_format="html",
        char_count=len(content),
        page_count=page_count,
    )


def _md_to_html(md_text: str) -> str:
    """Convert markdown (with tables) to HTML."""
    import markdown
    return markdown.markdown(md_text, extensions=["tables", "fenced_code"])


async def _parse_pdf_with_vl(content_bytes: bytes) -> tuple[str, int]:
    """Parse PDF using the VL Service pipeline (MinerU/Qwen).

    1. Save bytes to temp file
    2. Render pages to images (PyMuPDF)
    3. OCR each page with VL Service
    4. Convert each page markdown to HTML
    5. Return concatenated HTML + page count
    """
    import fitz
    from app.services.vl_service import vl_service

    # Write to temp file so fitz can open it
    tmp_path = Path(tempfile.mktemp(suffix=".pdf"))
    tmp_path.write_bytes(content_bytes)

    try:
        # Render pages to images
        doc = fitz.open(str(tmp_path))
        total_pages = len(doc)

        # Save page images to a temp directory
        img_dir = Path(tempfile.mkdtemp(prefix="temp_upload_"))
        image_paths: list[Path] = []

        for page_num in range(total_pages):
            page = doc[page_num]
            zoom = 3.0
            mat = fitz.Matrix(zoom, zoom)
            pix = page.get_pixmap(matrix=mat)
            img_path = img_dir / f"page_{page_num + 1}.png"
            pix.save(str(img_path))
            image_paths.append(img_path)

        doc.close()

        # OCR each page with VL Service → markdown → HTML
        html_parts: list[str] = []
        for i, img_path in enumerate(image_paths):
            page_num = i + 1
            try:
                md_content, _figures = await vl_service.ocr_page_to_markdown(img_path)
                page_html = _md_to_html(md_content)
                html_parts.append(
                    f'<section class="page" data-page="{page_num}">\n'
                    f'<h2>第 {page_num} 页</h2>\n'
                    f'{page_html}\n'
                    f'</section>'
                )
                logger.info(f"upload_temp: page {page_num}/{total_pages} OCR done")
            except Exception as e:
                logger.warning(f"upload_temp: page {page_num} OCR failed: {e}")
                html_parts.append(
                    f'<section class="page" data-page="{page_num}">\n'
                    f'<h2>第 {page_num} 页</h2>\n'
                    f'<p>[解析失败]</p>\n'
                    f'</section>'
                )

        content = "\n".join(html_parts)

        # Cleanup temp images
        shutil.rmtree(str(img_dir), ignore_errors=True)

        return content, total_pages

    finally:
        tmp_path.unlink(missing_ok=True)


def _parse_docx_to_html(content_bytes: bytes) -> tuple[str, int]:
    """Extract text from Word and convert to HTML. Returns (html, page_count)."""
    import markdown
    from docx import Document

    doc = Document(io.BytesIO(content_bytes))
    md_parts: list[str] = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        # Preserve heading levels
        style_name = para.style.name.lower() if para.style else ""
        if "heading 1" in style_name:
            md_parts.append(f"# {text}")
        elif "heading 2" in style_name:
            md_parts.append(f"## {text}")
        elif "heading 3" in style_name:
            md_parts.append(f"### {text}")
        else:
            md_parts.append(text)

    # Also extract tables
    for table in doc.tables:
        md_parts.append("")
        rows = table.rows
        if len(rows) == 0:
            continue
        # Header row
        header = "| " + " | ".join(cell.text.strip() for cell in rows[0].cells) + " |"
        separator = "| " + " | ".join("---" for _ in rows[0].cells) + " |"
        md_parts.append(header)
        md_parts.append(separator)
        # Data rows
        for row in rows[1:]:
            cells = "| " + " | ".join(cell.text.strip() for cell in row.cells) + " |"
            md_parts.append(cells)
        md_parts.append("")

    md_text = "\n".join(md_parts)
    html = markdown.markdown(md_text, extensions=["tables", "fenced_code"])

    page_count = max(1, len(md_text) // 3000)
    return html, page_count
