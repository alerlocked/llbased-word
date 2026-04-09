"""
Word文档导出服务
支持导出包含图片、注释、引用的完整稿件
"""
from pathlib import Path
from typing import List, Dict, Optional
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import re

from app.config import settings
from app.shared.logging import get_logger
logger = get_logger(__name__)
import urllib.parse


class WordExportService:
    """Word文档导出服务"""
    
    def __init__(self):
        """初始化服务"""
        self.export_dir = settings.DATA_DIR / "exports"
        self.export_dir.mkdir(parents=True, exist_ok=True)
    
    def create_document(
        self,
        title: str,
        content: str,
        annotations: Optional[List[Dict]] = None,
        citations: Optional[List[Dict]] = None,
        figures: Optional[List[Dict]] = None
    ) -> Document:
        """
        创建Word文档
        
        Args:
            title: 文章标题
            content: 文章内容（支持标记）
            annotations: 注释列表
            citations: 引用列表
            figures: 图片列表
        
        Returns:
            Document对象
        """
        logger.info(f"📄 创建Word文档: {title}")
        
        doc = Document()
        
        # 确保列表已初始化
        if annotations is None: annotations = []
        if citations is None: citations = []
        
        # 设置中文字体
        doc.styles['Normal'].font.name = '宋体'
        doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        doc.styles['Normal'].font.size = Pt(12)
        
        # 添加标题
        title_para = doc.add_heading(title, level=0)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 处理内容（解析图片、引用和注释标记）
        self._process_content(doc, content, figures, citations, annotations)
        
        # 添加注释部分
        if annotations and len(annotations) > 0:
            self._add_annotations(doc, annotations)
        
        # 添加参考文献
        if citations and len(citations) > 0:
            self._add_references(doc, citations)
        
        logger.info("✅ Word文档创建完成")
        return doc
    
    def _resolve_image_path(self, image_url: str) -> Optional[Path]:
        """
        解析图片URL，转换为本地文件路径
        
        Args:
            image_url: 图片URL（可能是相对路径、绝对路径或网络URL）
        
        Returns:
            本地文件路径，如果无法解析则返回None
        """
        if not image_url:
            return None
        
        # 处理相对路径（/static/data/...）
        if image_url.startswith('/static/data/'):
            # 去掉 /static/data/ 前缀，得到相对于 DATA_DIR 的路径
            relative_path = image_url.replace('/static/data/', '')
            img_path = settings.DATA_DIR / relative_path
            if img_path.exists():
                return img_path
            else:
                logger.warning(f"⚠️ 相对路径图片不存在: {img_path}")
                return None
        
        # 处理绝对路径
        if image_url.startswith('/'):
            img_path = Path(image_url)
            if img_path.exists():
                return img_path
        
        # 处理完整路径（包含盘符的Windows路径）
        if len(image_url) > 2 and image_url[1] == ':':
            img_path = Path(image_url)
            if img_path.exists():
                return img_path
        
        # 处理网络URL
        if image_url.startswith('http://') or image_url.startswith('https://'):
            logger.warning(f"⚠️ 网络图片URL无法直接导出到Word: {image_url}")
            # 可以尝试从WebImage表查找已下载的本地路径
            # 注意：这里需要传入db session，暂时跳过数据库查询
            # 如果需要在导出时查询数据库，应该在调用export_article时传入db session
            return None
        
        # 尝试作为相对路径处理（相对于DATA_DIR）
        img_path = settings.DATA_DIR / image_url
        if img_path.exists():
            return img_path
        
        # 尝试作为相对于BASE_DIR的路径
        img_path = settings.BASE_DIR / image_url
        if img_path.exists():
            return img_path
        
        logger.warning(f"⚠️ 无法解析图片路径: {image_url}")
        return None
    
    def _process_content(
        self,
        doc: Document,
        content: str,
        figures: Optional[List[Dict]] = None,
        citations: Optional[List[Dict]] = None,
        annotations: Optional[List[Dict]] = None
    ):
        """
        处理文章内容，解析图片、引用和注释标记
        
        Args:
            doc: Document对象
            content: 文章内容
            figures: 图片列表
            citations: 引用列表
            annotations: 注释列表
        """
        # 分割段落
        paragraphs = content.split('\n')
        
        figure_counter = 1
        citation_counter = 1
        annotation_counter = 1
        
        # 如果没有传入结构化数据，我们将从文本中提取
        extracted_citations = []
        extracted_annotations = []
        
        for para_text in paragraphs:
            para_text = para_text.strip()
            if not para_text:
                continue
            
            # 跳过"标题："行
            if para_text.startswith('标题：') or para_text.startswith('标题:'):
                continue
            
            # 1. 检查是否是Markdown图片语法：![alt](url)
            markdown_image_match = re.match(r'^!\[(.*?)\]\((.*?)\)$', para_text)
            if markdown_image_match:
                alt_text = markdown_image_match.group(1).strip()
                image_url = markdown_image_match.group(2).strip()
                
                # 处理图片URL
                img_path = self._resolve_image_path(image_url)
                
                if img_path and img_path.exists():
                    try:
                        # 添加图片
                        para = doc.add_paragraph()
                        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        run = para.add_run()
                        run.add_picture(str(img_path), width=Inches(5))
                        
                        # 添加图注
                        caption_text = alt_text if alt_text else f"图{figure_counter}"
                        caption = doc.add_paragraph(f"图{figure_counter}：{caption_text}")
                        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        try:
                            caption.style = 'Caption'
                        except:
                            pass
                        
                        figure_counter += 1
                        logger.info(f"✓ 插入Markdown图片: {alt_text} ({image_url})")
                    except Exception as e:
                        logger.error(f"❌ 插入Markdown图片失败: {str(e)}")
                        doc.add_paragraph(f"[图片占位：{alt_text}]")
                else:
                    logger.warning(f"⚠️ Markdown图片路径不存在: {image_url}")
                    doc.add_paragraph(f"[图片占位：{alt_text}]")
                continue
            
            # 2. 检查是否是旧格式的图片标记：[图片：描述|路径]
            image_match = re.match(r'^\[图片[：:](.*?)\|?(.*?)\]$', para_text)
            if image_match:
                description = image_match.group(1).strip()
                image_path_str = image_match.group(2).strip() if image_match.group(2) else ""
                
                # 尝试插入图片
                if image_path_str:
                    try:
                        # 处理路径，确保是绝对路径或相对于 BASE_DIR
                        img_path = Path(image_path_str)
                        if not img_path.is_absolute():
                            img_path = settings.BASE_DIR / image_path_str
                            
                        if img_path.exists():
                            # 添加图片
                            para = doc.add_paragraph()
                            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            run = para.add_run()
                            run.add_picture(str(img_path), width=Inches(5))
                            
                            # 添加图注
                            caption = doc.add_paragraph(f"图{figure_counter}：{description}")
                            caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            # 确保 Caption 样式存在
                            try:
                                caption.style = 'Caption'
                            except:
                                pass
                            
                            figure_counter += 1
                            logger.info(f"✓ 插入图片: {description}")
                        else:
                            logger.warning(f"⚠️ 图片不存在: {img_path}")
                            doc.add_paragraph(f"[图片占位：{description}]")
                    except Exception as e:
                        logger.error(f"❌ 插入图片失败: {str(e)}")
                        doc.add_paragraph(f"[图片占位：{description}]")
                else:
                    doc.add_paragraph(f"[图片占位：{description}]")
                continue

            # 3. 处理段落内的引用标记 [引用：来源]
            def replace_citation(match):
                nonlocal citation_counter
                source = match.group(1).strip()
                ref_num = citation_counter
                citation_counter += 1
                extracted_citations.append({"id": ref_num, "title": source})
                return f"[{ref_num}]"
            
            para_text = re.sub(r'\[引用[：:](.*?)\]', replace_citation, para_text)

            # 4. 处理段落内的注释标记 [注释：内容]
            def replace_annotation(match):
                nonlocal annotation_counter
                ann_content = match.group(1).strip()
                ann_num = annotation_counter
                annotation_counter += 1
                extracted_annotations.append({"id": ann_num, "content": ann_content, "annotation_type": "说明"})
                # 使用上标格式表示注释
                return f"[注{ann_num}]"
            
            para_text = re.sub(r'\[注释[：:](.*?)\]', replace_annotation, para_text)
            
            # 5. 处理段落内可能包含的Markdown图片语法
            # 如果段落中包含 ![alt](url)，将其提取出来单独处理
            markdown_images_inline = re.finditer(r'!\[(.*?)\]\((.*?)\)', para_text)
            image_positions = []
            for match in markdown_images_inline:
                alt_text = match.group(1).strip()
                image_url = match.group(2).strip()
                start_pos = match.start()
                end_pos = match.end()
                image_positions.append((start_pos, end_pos, alt_text, image_url))
            
            # 如果有内联图片，需要分段处理
            if image_positions:
                last_pos = 0
                para = doc.add_paragraph()
                para.paragraph_format.first_line_indent = Pt(24)
                para.paragraph_format.line_spacing = 1.5
                
                for start_pos, end_pos, alt_text, image_url in image_positions:
                    # 添加图片前的文本
                    if start_pos > last_pos:
                        text_before = para_text[last_pos:start_pos]
                        if text_before.strip():
                            para.add_run(text_before)
                    
                    # 处理图片
                    img_path = self._resolve_image_path(image_url)
                    if img_path and img_path.exists():
                        try:
                            # 先结束当前段落，添加图片段落
                            para = doc.add_paragraph()
                            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            run = para.add_run()
                            run.add_picture(str(img_path), width=Inches(5))
                            
                            caption_text = alt_text if alt_text else f"图{figure_counter}"
                            caption = doc.add_paragraph(f"图{figure_counter}：{caption_text}")
                            caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            try:
                                caption.style = 'Caption'
                            except:
                                pass
                            
                            figure_counter += 1
                            logger.info(f"✓ 插入内联Markdown图片: {alt_text}")
                        except Exception as e:
                            logger.error(f"❌ 插入内联图片失败: {str(e)}")
                            para.add_run(f"[图片：{alt_text}]")
                    else:
                        para.add_run(f"[图片：{alt_text}]")
                    
                    last_pos = end_pos
                
                # 添加剩余文本
                if last_pos < len(para_text):
                    text_after = para_text[last_pos:]
                    if text_after.strip():
                        para = doc.add_paragraph()
                        para.paragraph_format.first_line_indent = Pt(24)
                        para.paragraph_format.line_spacing = 1.5
                        para.add_run(text_after)
            else:
                # 没有内联图片，正常添加段落
                para = doc.add_paragraph(para_text)
                para.paragraph_format.first_line_indent = Pt(24)  # 首行缩进两字符
                para.paragraph_format.line_spacing = 1.5  # 1.5倍行距

        # 如果外部没有提供 annotations/citations，使用提取的结果
        if not annotations:
            annotations.extend(extracted_annotations)
        if not citations:
            citations.extend(extracted_citations)
    
    def _add_annotations(self, doc: Document, annotations: List[Dict]):
        """
        添加注释部分
        
        Args:
            doc: Document对象
            annotations: 注释列表
        """
        logger.info(f"📝 添加 {len(annotations)} 条注释")
        
        # 添加分页
        doc.add_page_break()
        
        # 添加注释标题
        doc.add_heading('注释', level=1)
        
        for idx, ann in enumerate(annotations, 1):
            content = ann.get('content', '')
            ann_type = ann.get('annotation_type', 'note')
            
            # 添加注释条目
            para = doc.add_paragraph()
            para.add_run(f"[{idx}] ").bold = True
            para.add_run(f"（{ann_type}）{content}")
    
    def _add_references(self, doc: Document, citations: List[Dict]):
        """
        添加参考文献
        
        Args:
            doc: Document对象
            citations: 引用列表
        """
        logger.info(f"📚 添加 {len(citations)} 条参考文献")
        
        # 添加分页
        doc.add_page_break()
        
        # 添加参考文献标题
        doc.add_heading('参考文献', level=1)
        
        for idx, citation in enumerate(citations, 1):
            source_type = citation.get('source_type', 'unknown')
            author = citation.get('author', '未知')
            title = citation.get('title', '未知标题')
            date = citation.get('date', '')
            url = citation.get('url', '')
            
            # 格式化引用
            ref_text = f"[{idx}] "

            if source_type == 'web':
                ref_text += f"{author}. {title}. {date}. {url}"
            elif source_type == 'book':
                ref_text += f"{author}. {title}. {date}."
            elif source_type == 'document':
                ref_text += f"文档：{title}. {author}. {date}."
            else:
                ref_text += f"{author}. {title}. {date}."
            
            doc.add_paragraph(ref_text)
    
    def export_to_file(
        self,
        document: Document,
        filename: str
    ) -> Path:
        """
        导出文档到文件
        
        Args:
            document: Document对象
            filename: 文件名（不含扩展名）
        
        Returns:
            文件路径
        """
        # 确保文件名安全
        safe_filename = re.sub(r'[<>:"/\\|?*]', '_', filename)
        file_path = self.export_dir / f"{safe_filename}.docx"
        
        # 保存文档
        document.save(str(file_path))
        
        logger.info(f"✅ Word文档已保存: {file_path}")
        return file_path
    
    def export_article(
        self,
        title: str,
        content: str,
        annotations: Optional[List[Dict]] = None,
        citations: Optional[List[Dict]] = None,
        figures: Optional[List[Dict]] = None,
        filename: Optional[str] = None
    ) -> Path:
        """
        导出完整文章为Word文档
        
        Args:
            title: 标题
            content: 内容
            annotations: 注释
            citations: 引用
            figures: 图片
            filename: 文件名（可选）
        
        Returns:
            导出的文件路径
        """
        logger.info(f"📤 导出文章: {title}")
        
        # 创建文档
        doc = self.create_document(
            title,
            content,
            annotations,
            citations,
            figures
        )
        
        # 导出文件
        if filename is None:
            filename = title
        
        file_path = self.export_to_file(doc, filename)
        
        return file_path


# 全局实例
_word_export_service = None

def get_word_export_service() -> WordExportService:
    """获取Word导出服务单例"""
    global _word_export_service
    if _word_export_service is None:
        _word_export_service = WordExportService()
    return _word_export_service

